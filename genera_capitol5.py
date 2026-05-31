"""
Genera el fitxer Capitol5_Streamlit.docx amb el text del capitol 5 del TFG.
Executa: python genera_capitol5.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estils globals ────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    p.add_run(text)
    return p

def code_inline(text):
    return f'"{text}"'

# ── TÍTOL ─────────────────────────────────────────────────────────────────────
doc.add_heading("Capítol 5: Anàlisi Visual mitjançant Streamlit", level=0)
doc.add_paragraph()

# ── 5.1 ──────────────────────────────────────────────────────────────────────
h1("5.1 Introducció a l'eina de visualització")
para(
    "Per fer accessible l'anàlisi del corpus de 66.621 tweets recollits sobre la xarxa de "
    "Rodalies de Catalunya, s'ha desenvolupat una aplicació web interactiva utilitzant "
    "Streamlit [1], un framework de codi obert per a Python que permet construir aplicacions "
    "de dades sense necessitat de coneixements de desenvolupament web. Streamlit ofereix "
    "components interactius —gràfics, mapes, filtres, taules— que es defineixen directament "
    "en codi Python i es renderitzen automàticament en un navegador."
)
para(
    "L'arquitectura de l'aplicació es basa en un únic fitxer Python (app.py) que es carrega "
    "al servidor web de Streamlit. Tota la lògica de càrrega, preprocessament i visualització "
    "es troba en aquest fitxer. Per evitar recarregar els datasets en cada interacció de "
    "l'usuari, s'utilitza el decorador @st.cache_data, que emmagatzema en memòria el resultat "
    "de les funcions de càrrega i l'invalida únicament quan canvia el fitxer d'origen "
    "(comprovant el timestamp de modificació). Aquesta decisió elimina la necessitat d'un "
    "servidor de base de dades independent."
)
para(
    "La navegació entre les set seccions de l'aplicació es gestiona mitjançant un st.radio "
    "a la barra lateral esquerra (sidebar), de manera que els filtres globals —línia, idioma, "
    "caràcter del tweet i rang de dates— romanen sempre visibles i s'apliquen simultàniament "
    "a totes les pàgines."
)

# ── 5.2 ──────────────────────────────────────────────────────────────────────
h1("5.2 Fonts de dades i preprocessament")
para("L'aplicació llegeix quatre fitxers locals en el moment d'iniciar-se:")

table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
headers = table.rows[0].cells
headers[0].text = "Fitxer"
headers[1].text = "Descripció"
for cell in headers:
    for run in cell.paragraphs[0].runs:
        run.bold = True

data = [
    ("classificacio_estacions_ubicacions.csv",
     "Dataset principal: 66.227 tweets amb estació detectada, línia associada, idioma "
     "i tipologia del tweet (informatiu / queixa / indefinit)"),
    ("stations_info.json",
     "Coordenades geogràfiques (latitud i longitud) de cada estació de les vuit línies, "
     "indexades per línia i per nom d'estació"),
    ("tweets_incidencias.csv",
     "Classificació de 66.621 tweets amb el tipus d'incident detectat (demora, averia, "
     "parada, etc.), la confiança del model i el mètode de classificació emprat"),
    ("tweets_merged.csv",
     "Dataset unificat generat mitjançant un join posicional de les tres fonts anteriors, "
     "que afegeix el nom d'usuari, el contingut original i les metadades d'incident a cada tweet"),
]
for i, (f, d) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = f
    row[1].text = d
doc.add_paragraph()

para("El preprocessament es realitza en quatre funcions cacheades:")
for item in [
    ("load_tweets()", "Llegeix el CSV principal, parseja l'estació associada (emmagatzemada "
     "en format JSON dins d'una columna), construeix les columnes de data, mes, hora i "
     "normalitza el valor 'mixt' a 'indefinit' per coherència terminològica."),
    ("build_expanded()", "Desnormalitza la relació tweet–estació: cada tweet que menciona N "
     "estacions genera N files al dataframe expandit, afegint les coordenades geogràfiques "
     "i la línia dominant per a cada estació."),
    ("load_incidents()", "Carrega el CSV d'incidències, desduplicant els tweets per id i "
     "mantenint la classificació del mètode més fiable (prioritat: llm_confirm > rules > cache)."),
    ("load_master()", "Carrega el dataset unificat tweets_merged.csv i construeix una columna "
     "datetime combinant data i hora."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

para(
    "\nLes dades es filtren al rang temporal 01/08/2025 – 30/06/2026, que correspon al "
    "període de recollida amb dades suficients per a l'anàlisi."
)

# ── 5.3 ──────────────────────────────────────────────────────────────────────
h1("5.3 Disseny visual i CSS personalitzat")
para(
    "Per donar a l'aplicació una aparença professional de dashboard de dades, s'ha injectat "
    "un bloc CSS personalitzat mitjançant st.markdown(..., unsafe_allow_html=True) "
    "immediatament després de st.set_page_config. Els principals elements de disseny són:"
)
for item in [
    ("Fons fosc", "el cos de la pàgina i la barra lateral utilitzen el color #080e1a "
     "(blau molt fosc), coherent amb l'estètica de les eines d'anàlisi de dades modernes."),
    ("Tipografia Inter", "s'importa des de Google Fonts i s'aplica globalment a tots els "
     "elements de text."),
    ("Metric cards", "les targetes de mètriques de Streamlit (st.metric) s'han estilitzat "
     "amb fons degradat, vora subtil i efecte d'elevació al passar el cursor (hover)."),
    ("Navegació de la barra lateral", "els botons de pàgina es mostren com a elements de "
     "llista sense vora, amb fons transparent i canvi de color en hover."),
    ("Headers de mapa", "cada secció de mapa utilitza un bloc HTML amb gradient lateral i "
     "barra de color d'accent que coincideix amb la línia o secció corresponent."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

# ── 5.4 ──────────────────────────────────────────────────────────────────────
h1("5.4 Filtres globals i navegació")
para(
    "La barra lateral conté dos blocs principals: la navegació (a dalt) i els filtres "
    "globals (a baix). La navegació utilitza st.radio amb set opcions i "
    "label_visibility=\"collapsed\" per ocultar l'etiqueta visible."
)
para("Els filtres globals —disponibles a totes les pàgines— inclouen:")
for item in [
    ("Línia (st.pills, selecció múltiple)", "R1, R2, R2N, R2S, R3, R4, R7, R8."),
    ("Idioma (st.pills, selecció múltiple)", "català, castellà, anglès, indeterminat."),
    ("Caràcter (st.pills, selecció múltiple)", "informatiu, queixa, indefinit."),
    ("Període de temps (st.date_input amb rang)",
     "rang seleccionable entre 01/08/2025 i 30/06/2026, amb un botó ↺ adjacent "
     "que restableix el rang complet en un sol clic."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

para(
    "\nLa funció apply_filters() centralitza tota la lògica de filtratge. Quan "
    "use_lines_list=True, utilitza una expressió regular per fer coincidir la columna "
    "lines_list (format pipe-separated) evitant falsos positius entre línies de nom "
    "similar (p. ex. R2 vs. R2N)."
)

# ── 5.5 ──────────────────────────────────────────────────────────────────────
h1("5.5 Pàgines de l'aplicació")

h2("5.5.1 Pàgina d'inici (Inici)")
para(
    "La pàgina d'inici actua com a portal d'entrada al projecte. Mostra el títol "
    "\"Rodalies de Catalunya\" amb el subtítol que descriu l'objectiu del TFG, un badge "
    "amb la titulació i el crèdit d'autoria. A continuació es presenten vuit mètriques "
    "dividides en dos grups: les del dataset de tweets (total de tweets, tweets amb estació "
    "detectada, estacions identificades, dies amb dades) i les de les incidències detectades "
    "(tweets d'incident, tipus més frequent, mesos amb incidents, percentatge sobre el total). "
    "Per acabar, dos gràfics de previsualització resumeixen la distribució dels incidents "
    "per tipus i l'evolució mensual."
)

h2("5.5.2 Mapa geogràfic")
para(
    "Aquesta pàgina és el nucli de la visualització espacial. Consta de quatre mapes "
    "apilats verticalment, tots ells actualitzats simultàniament en canviar qualsevol "
    "filtre de la barra lateral:"
)
for item in [
    ("Mapa de punts per estació",
     "utilitza folium.CircleMarker amb radi proporcional al nombre de tweets per estació "
     "i folium.PolyLine per traçar les rutes de cada línia. Cada cercle té un popup "
     "enriquit amb les últimes cinc publicacions de l'estació. Les etiquetes d'estació "
     "es renderitzen amb folium.DivIcon i es poden filtrar per mode: tots els noms, "
     "top N estacions o codis de línia."),
    ("Mapa de calor",
     "empra folium.plugins.HeatMap amb els valors de densitat de tweets per estació "
     "com a pes. S'hi superposen les mateixes etiquetes d'estació que al mapa de punts."),
    ("Mapa 3D per zones",
     "implementat amb pydeck.Layer(\"HexagonLayer\") sobre un basemap de CartoDB Dark Matter. "
     "Els hexàgons s'extrueixen verticalment en proporció a la densitat de tweets; el radi "
     "de 3.000 metres assegura que les zones es fusionin visualment. Una capa addicional "
     "ScatterplotLayer superposa les estacions individuals amb tooltip que mostra el nom "
     "i el recompte de tweets."),
    ("Mapa de clústers",
     "utilitza folium.plugins.MarkerCluster per agrupar les estacions i mostrar el nombre "
     "total en cada clúster, facilitant la identificació visual de les zones de major activitat."),
]:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

h2("5.5.3 Anàlisi temporal")
para(
    "La pàgina d'anàlisi temporal examina l'evolució al llarg del temps. Incorpora un "
    "filtre Top N que, quan s'activa, restringeix tots els gràfics als N dies amb més "
    "activitat del període seleccionat. Les visualitzacions en ordre:"
)
for item in [
    ("Gràfic de sectors (donut)", "distribució dels tweets per tipologia del període. "
     "El nombre total es mostra al centre de l'anell. A sota, cards individuals mostren "
     "exemples de tweets de cada categoria."),
    ("Volum diari", "gràfic de línies amb marcadors. Els N dies més actius es destaquen "
     "amb una estrella vermella."),
    ("Top N dies", "per a cada dia més actiu, una card HTML mostra la data, el nombre "
     "de tweets i fins a tres exemples del contingut publicat."),
    ("Comparativa mensual", "gràfic de barres verticals de tots els mesos del període. "
     "Els mesos amb els dies del Top N es pinten de vermell; la resta, de blau."),
    ("Distribució horària", "barres apilades per hora del dia (0–23) amb línia de total "
     "superposa. S'alimenta del dataset sense el filtre Top N."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

h2("5.5.4 Anàlisi per línies")
para(
    "Aquesta pàgina permet comparar l'activitat entre línies i analitzar el comportament "
    "per estació. Un filtre Top N opcional limita la visualització a les N línies amb més "
    "tweets. El primer gràfic és un diagrama de barres horitzontal que ordena les línies "
    "per nombre de tweets únics. Un selector de línia (st.pills) permet aprofundir en "
    "qualsevol línia i veure, per a cada estació, la distribució percentual de tweets per "
    "caràcter en un gràfic de barres apilades horitzontal."
)

h2("5.5.5 Anàlisi d'incidències")
para(
    "La pàgina d'incidències treballa exclusivament amb el dataset tweets_incidencias.csv. "
    "Incorpora dos filtres propis —tipus d'incident i nivell de confiança del classificador— "
    "a més dels filtres de data globals. Les visualitzacions inclouen: mètriques globals, "
    "un gràfic de barres apilades diàries per tipus d'incident, un gràfic horitzontal de "
    "distribució total per tipus, i un expander amb la taula completa de tweets."
)

h2("5.5.6 Incidències per línia")
para(
    "Aquesta pàgina creua les dues fonts de dades principals: el dataset d'incidències i "
    "el d'estacions, units per identificador de tweet. Per a cada línia activa es genera "
    "un gràfic de barres horitzontal amb el recompte de tweets per tipus d'incident. "
    "En clicar sobre una barra (on_select=\"rerun\"), apareix una llista de tweets d'aquell "
    "tipus amb data, hora, confiança i text. Un filtre d'hora local (toggle + slider 0–23) "
    "permet restringir l'anàlisi a tweets publicats a partir d'una determinada hora."
)

h2("5.5.7 Cas d'Estudi: 20 de Gener de 2026")
para(
    "Aquesta pàgina representa l'aplicació pràctica central del TFG: la demostració que "
    "les publicacions dels usuaris de Twitter permeten detectar incidents ferroviaris amb "
    "antelació respecte a les comunicacions oficials de Rodalies."
)
para(
    "El 20 de gener de 2026 va ser el dia amb el major volum de tweets de tot el dataset "
    "(1.255 publicacions), a causa de múltiples incidents simultanis provocats per un "
    "temporal: una caiguda d'arbre a les vies de la línia R11 entre Breda i "
    "Maçanet-Massanes, diverses interrupcions a la línia R4 i avaries d'infraestructura "
    "a les línies R2 i R1."
)
para("La pàgina s'estructura en tres blocs:")
for item in [
    ("Resum per línia",
     "per a cada línia activa, es mostren el gràfic de distribució de tipus d'incident i, "
     "en paral·lel, els tweets d'usuaris publicats abans del primer avís oficial de "
     "@rodalies i el primer avís oficial identificat. Un badge de color calcula i mostra "
     "el temps d'avantatge: vermell si supera 60 minuts, taronja entre 30 i 60 minuts, "
     "verd si és inferior a 30 minuts."),
    ("Bloc especial R11",
     "a les 08:05, l'usuari @korexpan va publicar: \"Caiguda d'arbre a la via i s'anulen "
     "R11 @rodalies Figueres - Sants\", identificant correctament la causa de l'incident. "
     "La primera comunicació oficial de @rodalies mencionant explícitament la caiguda "
     "d'un arbre no va arribar fins a les 10:26, 2 hores i 21 minuts més tard."),
    ("Tweets sense línia assignada",
     "recull els tweets classificats com a incident (confiança > 0,80) que no han pogut "
     "ser associats a cap línia concreta."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

# ── 5.6 ──────────────────────────────────────────────────────────────────────
h1("5.6 Llibreries principals")

table2 = doc.add_table(rows=8, cols=3)
table2.style = "Table Grid"
headers2 = table2.rows[0].cells
for i, h in enumerate(["Llibreria", "Versió", "Ús a l'aplicació"]):
    headers2[i].text = h
    for run in headers2[i].paragraphs[0].runs:
        run.bold = True

libs = [
    ("Streamlit [1]", "1.52.1",
     "Framework principal: interfície web, widgets, cache, navegació"),
    ("Pandas", "2.x",
     "Càrrega, transformació i filtratge dels datasets"),
    ("Folium [2]", "0.18.x",
     "Mapes interactius de Leaflet.js: punts, polilínies, heatmap, clústers"),
    ("streamlit-folium", "0.23.x",
     "Integració dels mapes Folium dins de components Streamlit"),
    ("Plotly [3]", "5.x",
     "Gràfics interactius: barres, sectors, línies, barres apilades"),
    ("PyDeck [4]", "0.9.x",
     "Mapa 3D WebGL amb capa HexagonLayer i ScatterplotLayer"),
    ("Folium.plugins", "—",
     "HeatMap i MarkerCluster com a capes addicionals de Folium"),
]
for i, (lib, ver, use) in enumerate(libs, 1):
    row = table2.rows[i].cells
    row[0].text = lib
    row[1].text = ver
    row[2].text = use
doc.add_paragraph()

# ── 5.7 ──────────────────────────────────────────────────────────────────────
h1("5.7 Limitacions tècniques")
for item in [
    ("Interactivitat dels mapes Folium",
     "Streamlit no disposa d'un mecanisme natiu per capturar clics sobre elements de "
     "Folium. La solució implementada per als gràfics de Plotly (on_select=\"rerun\") "
     "no és aplicable als mapes de Leaflet, de manera que les interaccions als mapes "
     "es limiten als tooltips i popups natius de Folium."),
    ("Invalidació de la caché",
     "El decorador @st.cache_data utilitza el timestamp de modificació del fitxer CSV "
     "com a clau de caché. Si les dades s'actualitzen sense reiniciar l'aplicació, "
     "la caché no s'invalida automàticament."),
    ("Execució local",
     "L'aplicació s'executa en local (http://localhost:8503) i no ha estat desplegada "
     "a Streamlit Cloud ni a cap servidor públic, ja que el dataset conté tweets de "
     "tercers i el desplegament públic requeriria una revisió de la política de privacitat."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item[0])
    r.bold = True
    p.add_run(f": {item[1]}")

# ── REFERÈNCIES ───────────────────────────────────────────────────────────────
doc.add_paragraph()
h1("Referències del capítol")
for ref in [
    "[1] Streamlit Inc. Streamlit documentation. Disponible a: https://docs.streamlit.io (consultat: juny 2026)",
    "[2] Python-Visualization. Folium documentation. Disponible a: https://python-visualization.github.io/folium/ (consultat: juny 2026)",
    "[3] Plotly Technologies Inc. Plotly Python Open Source Graphing Library. Disponible a: https://plotly.com/python/ (consultat: juny 2026)",
    "[4] vis.gl. deck.gl / pydeck documentation. Disponible a: https://deckgl.readthedocs.io (consultat: juny 2026)",
]:
    doc.add_paragraph(ref, style="List Number")

# ── GUARDAR ───────────────────────────────────────────────────────────────────
out = r"C:\MARINA\Universitat\TFG - Visualització\Capitol5_Streamlit.docx"
doc.save(out)
print(f"Guardat: {out}")

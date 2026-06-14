"""
Genera el document Word del Capitol 6 — Resultats del TFG de Marina Castellano.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estils globals ─────────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(6)
    return h

def add_para(text, italic=False, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.italic = italic
    run.font.bold   = bold
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"[Nota per a la Marina: {text}]")
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(11)
    run.font.bold   = True
    run.font.italic = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Capçalera
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
    # Files
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    doc.add_paragraph()  # espai despres de la taula
    return t


# ══════════════════════════════════════════════════════════════════════════════
# CAPITOL 6
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. Resultats", level=1)

# ── 6.1 ──────────────────────────────────────────────────────────────────────
add_heading("6.1 Visió general dels resultats obtinguts", level=2)

add_para(
    "El sistema desenvolupat al llarg d'aquest treball ha produït un conjunt de resultats que "
    "cobreixen quatre dimensions complementàries: la recollida i preparació d'un corpus de "
    "publicacions a Twitter/X, la classificació automàtica del seu contingut, l'extracció "
    "d'informació geogràfica i operativa, i l'avaluació de la capacitat de detecció anticipada "
    "d'incidències. En els apartats que segueixen s'analitza cadascuna d'aquestes dimensions de "
    "manera individual, i el capítol es tanca amb una revisió del grau de compliment dels "
    "objectius específics formulats a la Secció 1.3."
)
add_para(
    "En conjunt, el corpus recollit conté 66.621 tweets sobre la xarxa de Rodalies de Catalunya, "
    "classificats amb una confiança mitjana de 0,8975, dels quals 8.282 fan referència a incidents "
    "operatius reals. L'anàlisi dels quatre casos d'estudi mostra que, en situacions d'incidència "
    "sobtada, les publicacions dels usuaris precedeixen l'avís oficial de Rodalies en intervals "
    "que van dels 5 minuts fins a les 2 hores i 21 minuts, en funció de la naturalesa i l'escala "
    "de l'incident."
)

# ── 6.2 ──────────────────────────────────────────────────────────────────────
add_heading("6.2 Resultats de la recollida i preparació de dades", level=2)

add_para(
    "Tal com es descriu a la Secció 3.3, el procés de scraping va recollir publicacions relatives "
    "a les vuit línies de Rodalies de Catalunya des del 1 d'agost de 2025 fins al 30 de juny de "
    "2026. El corpus final, un cop aplicada la deduplicació i el processament descrit a la "
    "Secció 3.4, conté 66.621 tweets distribuïts en 261 dies amb activitat registrada. La mediana "
    "de publicacions per dia és de 187,0, amb una mitjana de 253,4; la dispersió per sobre de la "
    "mitjana respon principalment als dies d'incidència greu, que generen volums molt superiors "
    "als habituals."
)
add_para(
    "Quant a la distribució lingüística, el 82,6 % dels tweets estan escrits en català (55.040 "
    "publicacions), el 9,4 % en castellà (6.248), el 7,0 % en idioma indeterminat (4.660) i el "
    "0,4 % en anglès (279). Aquesta distribució reflecteix la base d'usuaris del servei i és "
    "coherent amb el perfil sociolingüístic de l'àrea metropolitana de Barcelona i les comarques "
    "que connecten les vuit línies."
)
add_para(
    "La classificació del caràcter de cadascun dels tweets revela que tres de cada quatre "
    "publicacions (74,8 %, n = 49.854) es corresponen a tweets informatius —consultes, "
    "actualitzacions d'estat o transmissió de contingut oficial—, mentre que el 20,2 % "
    "(n = 13.442) expressen queixes o valoracions negatives sobre el servei. El 4,4 % restant "
    "(n = 2.931) presenta un caràcter mixt i ha estat normalitzat a la categoria «indefinit» "
    "en el processament, tal com s'explica a la Secció 4.2.4."
)
add_para(
    "La Taula 6.1 recull la distribució del corpus pels 12.629 tweets als quals el sistema ha "
    "pogut assignar una línia de Rodalies concreta. Cal tenir en compte que aquesta xifra "
    "representa el 19,0 % del total, ja que la majoria dels usuaris no incorporen el codi de "
    "línia de manera explícita en les seves publicacions. El mecanisme de detecció d'estacions "
    "i línies es descriu detalladament a la Secció 4.2."
)

add_table_caption("Taula 6.1. Distribució del corpus per línia de Rodalies (tweets amb línia assignada, n = 12.629).")
add_table(
    ["Línia", "Tweets", "%"],
    [
        ["R4",  "3.266", "25,9 %"],
        ["R1",  "3.204", "25,4 %"],
        ["R2S", "1.896", "15,0 %"],
        ["R2",  "1.450", "11,5 %"],
        ["R2N", "1.361", "10,8 %"],
        ["R3",    "929",  "7,4 %"],
        ["R8",    "432",  "3,4 %"],
        ["R7",     "91",  "0,7 %"],
    ]
)

add_para(
    "Les línies R4 i R1 concentren gairebé la meitat del corpus amb línia assignada (51,3 %), "
    "fet consistent amb la seva major demanda de viatgers i la cobertura de les àrees "
    "metropolitanes de Baix Llobregat-Anoia i Maresme-Vallès, respectivament."
)

# ── 6.3 ──────────────────────────────────────────────────────────────────────
add_heading("6.3 Resultats de la classificació automàtica d'incidències", level=2)

add_para(
    "El sistema híbrid de classificació descrit a la Secció 4.3 ha processat els 66.621 tweets "
    "del corpus produint una etiqueta de tipus d'incident i una puntuació de confiança per a "
    "cada publicació. En la resolució de conflictes, s'ha aplicat l'ordre de prioritat: "
    "llm_confirm > rules > cache, de manera que el mètode rules —que aplica el conjunt de "
    "regles deterministes— ha resolt el 84,0 % dels casos, el model LLM ha intervingut com a "
    "àrbitre en el 7,8 % i el sistema de memòria cau ha cobert el 7,6 % restant (Taula 6.2). "
    "Un 0,6 % de tweets (n = 394) no han pogut ser classificats per cap dels tres mètodes i "
    "han quedat amb etiqueta buida."
)

add_table_caption("Taula 6.2. Distribució del corpus per mètode de classificació.")
add_table(
    ["Mètode", "Tweets", "%"],
    [
        ["rules",       "55.966", "84,0 %"],
        ["llm_confirm",  "5.207",  "7,8 %"],
        ["cache",        "5.054",  "7,6 %"],
        ["(buit)",         "394",  "0,6 %"],
    ]
)

add_para(
    "La confiança mitjana del conjunt és de 0,8975, i el 98,4 % dels tweets (n = 65.550) "
    "superen el llindar de 0,80, que és el valor a partir del qual el sistema considera que "
    "la classificació és suficientment fiable per a l'anàlisi comparativa dels casos d'estudi. "
    "Únicament 547 tweets (0,8 %) han estat marcats com a «en dubte» durant el procés de "
    "classificació manual d'arbitratge."
)
add_para(
    "Dels 66.621 tweets classificats, 8.282 corresponen a incidents operatius reals (és a dir, "
    "excloent la categoria sin_incidencia), fet que representa el 12,4 % del corpus. La Taula "
    "6.3 detalla la distribució per tipus d'incident: les demores constitueixen el tipus "
    "majoritari (51,7 % dels incidents reals), seguides de les obres en via (16,6 %) i les "
    "parades no programades (15,5 %). Els incidents d'arrollament, per la seva excepcionalitat, "
    "representen el percentatge més baix (1,8 %)."
)

add_table_caption("Taula 6.3. Distribució dels tweets classificats com a incidents reals per tipus (n = 8.282).")
add_table(
    ["Tipus d'incident", "Tweets", "%"],
    [
        ["demora",        "4.279", "51,7 %"],
        ["obras",         "1.376", "16,6 %"],
        ["parada",        "1.281", "15,5 %"],
        ["huelga",          "832", "10,0 %"],
        ["averia",          "362",  "4,4 %"],
        ["arrollamiento",   "152",  "1,8 %"],
    ]
)

add_para(
    "El predomini de les demores és esperable en un servei ferroviari de rodalia, on qualsevol "
    "incident de durada inferior als 30 minuts es manifesta principalment com a retard percebut "
    "pels usuaris. Les obres i parades reflecteixen el context del període analitzat, marcat per "
    "les actuacions de manteniment i reparació de la infraestructura posteriors al temporal de "
    "gener de 2026."
)
add_para(
    "Cal precisar que l'eina de visualització interactiva (Capítol 5) aplica un filtre "
    "temporal que restringeix la vista al període principal d'anàlisi (agost 2025 – juny "
    "2026), excloent 488 tweets amb dates anteriors a aquesta finestra que formen part del "
    "corpus classificat però no del tauler interactiu. Per aquesta raó, el panell d'inici "
    "de l'aplicació mostra 7.888 tweets d'incident, mentre que la xifra del corpus complet "
    "és de 8.282."
)

# ── 6.4 ──────────────────────────────────────────────────────────────────────
add_heading("6.4 Resultats de l'extracció d'estacions i línies", level=2)

add_para(
    "El mòdul de reconeixement d'entitats geogràfiques descrit a la Secció 4.2 ha processat els "
    "66.621 tweets del corpus i ha identificat mencions d'estacions de la xarxa en 66.227 d'ells. "
    "D'aquests, tan sols 12.629 han pogut rebre una assignació de línia definitiva, la qual cosa "
    "suposa una taxa de cobertura del 19,0 %. Aquesta xifra és inherent al fenomen analitzat: els "
    "usuaris de Twitter/X rarament incorporen el codi de línia de forma explícita; quan descriuen "
    "una incidència tendeixen a esmentar l'estació o la destinació, i en molts casos ni tan sols "
    "ho fan."
)
add_para(
    "El sistema, tal com s'explica a la Secció 4.2.3, prioritza la precisió per sobre del recall: "
    "en situacions ambigües no s'assigna línia per evitar associacions incorrectes que podrien "
    "introduir soroll en l'anàlisi comparativa. Aquesta decisió de disseny implica acceptar que "
    "una fracció significativa del corpus queda sense informació geogràfica assignada, però "
    "garanteix que els tweets amb línia identificada siguen fiables. La validació geogràfica dels "
    "casos d'estudi del present capítol es basa exclusivament en els tweets amb assignació "
    "confirmada."
)

# ── 6.5 ──────────────────────────────────────────────────────────────────────
add_heading("6.5 Resultats de l'anàlisi temporal i geogràfica", level=2)

add_para(
    "L'anàlisi de la distribució temporal del corpus revela un patró diari consistent amb les "
    "franges d'utilització del servei de Rodalies: la concentració de publicacions és màxima "
    "entre les 7 h i les 10 h del matí i entre les 17 h i les 20 h del vespre, coincidint "
    "amb les hores punta d'entrada i sortida de la feina. Fora d'aquestes franges, l'activitat "
    "disminueix de manera notable, amb un mínim en les hores centrals de la nit."
)
add_para(
    "Al llarg del període analitzat, la setmana del 20 al 27 de gener de 2026 concentra el "
    "major volum d'activitat del dataset: els efectes del temporal hivernal i les seves seqüeles "
    "en la infraestructura ferroviària van generar una activitat sostinguda durant vuit dies "
    "consecutius per sobre dels 800 tweets diaris. El dia amb el pic màxim absolut va ser el "
    "26 de gener (1.269 tweets), cinc dies després de l'inici de les afectacions. El 13 de "
    "febrer de 2026, coincidint amb la Borrasca Nils, es va registrar el valor màxim puntual "
    "fora d'aquell episodi (1.002 tweets). En contrast, en dies sense incidències rellevants, "
    "el corpus se situa per sota de la mediana de 187,0 tweets diaris."
)
add_para(
    "Des del punt de vista geogràfic, les estacions del nucli de Barcelona (Sants, Passeig de "
    "Gràcia, Arc de Triomf) concentren el major volum de tweets al llarg de tot el període, de "
    "manera consistent amb el seu rol de nodes centrals de la xarxa. No obstant, en els dies "
    "dels casos d'estudi s'observen concentracions puntuals a estacions perifèriques directament "
    "afectades per l'incident: Breda i Maçanet-Massanes durant el temporal de gener, La Garriga "
    "i Figaró durant el robatori de coure de març, o les estacions tancades durant la Borrasca "
    "Nils (Malgrat de Mar, Premià de Mar, Barberà del Vallès). La visualització interactiva "
    "desenvolupada (vegeu Capítol 5) permet observar aquesta dinàmica espacial de forma directa "
    "als mapes de calor i de punts."
)
add_para(
    "Les línies R4 i R1, que concentren el major volum de tweets, corresponen a les línies amb "
    "major nombre de viatgers diaris de la xarxa i cobreixen les zones de major activitat laboral "
    "a l'entorn de Barcelona. Les línies R7 i R8, menys esmentades al corpus, presenten volums "
    "d'activitat social molt inferiors, en línia amb la seva menor demanda i la seva cobertura "
    "de zones menys densament poblades."
)

# ── 6.6 ──────────────────────────────────────────────────────────────────────
add_heading("6.6 Cas d'estudi: detecció primerenca respecte als avisos oficials", level=2)

add_para(
    "Per avaluar la hipòtesi central del treball, s'han analitzat quatre dies amb incidències "
    "documentades a la xarxa de Rodalies de Catalunya. Per a cada línia activa en cada dia, "
    "s'ha identificat el primer tweet d'un usuari no oficial que el sistema classifiqui com a "
    "incident (tipus diferent de sin_incidencia) o queixa, i s'ha comparat la seva hora de "
    "publicació amb la del primer avís oficial d'un compte acreditat de Rodalies —@rodalies, "
    "@rod*cat, @renfe, o fonts informatives de referència com @3CatInfoViari— que contingui "
    "paraules clau d'incident (retard, tall, interrupció, demora, caiguda, temporal, afectació, "
    "suprimit, alternatiu). L'avantatge temporal és la diferència, expressada en minuts, entre "
    "ambdues hores. En els casos en què no s'ha identificat cap avís oficial dins la finestra "
    "del dia analitzat, la línia es classifica com a «senyal social sense avís oficial associat»."
)

# 6.6.1
add_heading("6.6.1 — 23 d'Octubre de 2025", level=3)

add_para(
    "El 23 d'octubre de 2025 es van registrar 506 publicacions, distribuïdes entre les vuit "
    "línies de la xarxa. La Taula 6.4 recull la comparativa entre el primer tweet d'usuari i "
    "el primer avís oficial per a cada línia activa."
)

add_table_caption("Taula 6.4. Detecció anticipada per línia, 23 d'Octubre de 2025.")
add_table(
    ["Línia", "Primer tweet usuari", "Primer avís oficial", "Avantatge", "n pre", "Estat"],
    [
        ["R4",  "07:36", "14:27", "+6 h 51 min", "5", "Detecció precoç"],
        ["R3",  "08:28", "10:37", "+2 h 9 min",  "1", "Detecció precoç"],
        ["R1",  "—",     "06:30", "—",            "0", "Sense detecció prèvia"],
        ["R7",  "—",     "12:38", "—",            "0", "Sense detecció prèvia"],
        ["R8",  "—",     "14:27", "—",            "0", "Sense detecció prèvia"],
        ["R2, R2N, R2S", "—", "—", "—",           "0", "Sense avís oficial associat"],
    ]
)

add_para(
    "Per a la línia R4, cinc tweets d'usuari publicats entre les 7:36 i les 14:27 precedeixen "
    "el primer avís oficial de gairebé set hores. Per a la línia R3, un sol tweet a les 8:28 "
    "anticipa la comunicació oficial de les 10:37 en 129 minuts. La mida mostral pre-oficial "
    "és reduïda en ambdós casos, de manera que l'avantatge temporal s'interpreta com a evidència "
    "qualitativa i no permet un contrast estadístic robust. Les línies R2, R2N i R2S no "
    "presenten avís oficial identificat dins la jornada, fet que pot respondre a incidències "
    "de menor magnitud no comunicades per canals de Twitter/X o a afectacions resoltes per "
    "altres vies."
)

# 6.6.2
add_heading("6.6.2 — 20 de Gener de 2026", level=3)

add_para(
    "El 20 de gener de 2026, un temporal hivernal va provocar afectacions simultànies a set "
    "línies de la xarxa, amb 807 publicacions registrades. L'activitat es va mantenir elevada "
    "durant la setmana posterior —el pic màxim del dataset es va produir el 26 de gener, amb "
    "1.269 tweets—, evidenciant que les disrupcions no es van resoldre de manera immediata."
)

# Paràgraf amb cites textuals dels tweets (en cursiva)
p_r11 = doc.add_paragraph()
p_r11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_r11.paragraph_format.first_line_indent = Inches(0.5)
p_r11.paragraph_format.space_after = Pt(6)
p_r11.add_run(
    "L'incident més rellevant va ser la caiguda d'un arbre a les vies de la línia R11, entre "
    "Breda i Maçanet-Massanes. A les 7:57 h, l'usuari @korexpan alertava: "
).font.name = "Times New Roman"
p_r11.runs[0].font.size = Pt(12)
cita1 = p_r11.add_run("«@rodalies us heu petat l’R11 de les 6.48 a Sant Celoni i aquí ningú informa»")
cita1.italic = True; cita1.font.name = "Times New Roman"; cita1.font.size = Pt(12)
p_r11.add_run(
    ". A les 8:05 h identificava la causa: "
).font.name = "Times New Roman"
p_r11.runs[2].font.size = Pt(12)
cita2 = p_r11.add_run("«Caiguda d’arbre a la via i s’anulen R11 @rodalies Figueres - Sants»")
cita2.italic = True; cita2.font.name = "Times New Roman"; cita2.font.size = Pt(12)
p_r11.add_run(
    ". El compte @3CatInfoViari va confirmar la interrupció a les 10:05 h i el compte oficial "
    "de @rodalies no va publicar la comunicació fins a les 10:26 h, amb un avantatge de "
).font.name = "Times New Roman"
p_r11.runs[4].font.size = Pt(12)
bold_141 = p_r11.add_run("141 minuts")
bold_141.bold = True; bold_141.font.name = "Times New Roman"; bold_141.font.size = Pt(12)
p_r11.add_run(
    " respecte al primer tweet d’identificació de la causa."
).font.name = "Times New Roman"
p_r11.runs[6].font.size = Pt(12)

add_para(
    "La Taula 6.5 recull la comparativa per al conjunt de línies. Quatre de les set presenten "
    "activitat social prèvia a l'avís oficial (R2N +81 min, R2S +74 min, R2 +33 min), mentre "
    "que R8 registra sis tweets d'usuari des de les 7:55 h sense comunicació oficial fins a les "
    "20:15 h. Les línies R4 i R3 van rebre avisos oficials abans de les 7:20 h, en un moment "
    "sense activitat social prèvia classificada, la qual cosa indica que en aquells casos el "
    "canal institucional va actuar amb rapidesa."
)

add_table_caption("Taula 6.5. Detecció anticipada per línia, 20 de Gener de 2026.")
add_table(
    ["Línia", "Primer tweet usuari", "Primer avís oficial", "Avantatge", "n pre", "Estat"],
    [
        ["R1/R11", "08:05", "10:26", "+2 h 21 min",  "4", "Detecció precoç"],
        ["R2N",    "07:57", "09:18", "+1 h 21 min",  "2", "Detecció precoç"],
        ["R2S",    "09:08", "10:22", "+1 h 14 min",  "3", "Detecció precoç"],
        ["R2",     "08:46", "09:19", "+33 min",       "1", "Detecció precoç"],
        ["R8",     "07:55", "20:15", "+12 h 20 min",  "6", "Senyal sense avís immediat"],
        ["R4",     "—",     "07:20", "—",             "0", "Sense detecció prèvia"],
        ["R3",     "—",     "07:19", "—",             "0", "Sense detecció prèvia"],
    ]
)

# 6.6.3
add_heading("6.6.3 — 13 de Febrer de 2026 (Borrasca Nils)", level=3)

add_para(
    "El 13 de febrer de 2026, la Borrasca Nils —considerada per l'Agència Estatal de "
    "Meteorologia una de les ventades més intenses de la dècada a Catalunya, amb ratxes "
    "superiors a 100 km/h— va provocar una afectació generalitzada a tota la xarxa de "
    "Rodalies. El dia va registrar 1.002 tweets, el valor màxim del dataset fora de la setmana "
    "posterior al temporal de gener. La combinació amb les seqüeles del descarrilament de "
    "Gelida (aproximadament 200 limitacions de velocitat actives i 71 punts d'obra en "
    "execució simultània) va originar retards mitjans superiors a 30 minuts, el tancament de "
    "les estacions de Malgrat de Mar, Premià de Mar i Barberà del Vallès, i la suspensió del "
    "servei de passatgers als trams R4 Sant Sadurní d'Anoia – Martorell, R3 i R8."
)
add_para(
    "A diferència dels casos anteriors, cap de les vuit línies actives aquell dia va presentar "
    "tweets d'usuari classificats com a incident o queixa anteriors al primer avís oficial. "
    "Tots els comptes de Rodalies van comunicar les afectacions entre les 07:15 i les 07:30 "
    "del matí, molt abans que els usuaris comencessin a publicar sobre els problemes específics "
    "de cada línia. Aquest resultat no invalida la hipòtesi del treball, sinó que la matisa: "
    "en events de gran escala previsibles amb antelació gràcies a les previsions meteorològiques, "
    "el canal oficial pot anticipar-se a la detecció social. Pender et al. (2014b) documenten "
    "que els usuaris de Twitter sovint detecten disrupcions ferroviàries abans que l'operadora "
    "hagi publicat cap avís oficial, però apunten que els events d'afectació massiva planificats "
    "o previsibles constitueixen un cas particular en el qual el canal institucional pot actuar "
    "amb rapidesa gràcies a l'anticipació operativa [Pender et al., 2014b]. En el cas de la "
    "Borrasca Nils, el servei meteorològic havia emès avisos de vent extrem amb hores "
    "d'antelació, la qual cosa permet inferir que Rodalies va activar els protocols de "
    "comunicació preventiva abans que les disrupcions es fessin perceptibles als viatgers."
)

# 6.6.4
add_heading("6.6.4 — 6 de Març de 2026", level=3)

add_para(
    "El 6 de març de 2026 es van registrar 774 tweets i l'incident de referència va ser un "
    "robatori de coure a La Garriga que va provocar retards d'uns 30 minuts a la línia R3, "
    "amb trens truncats a Figaró. La Taula 6.6 recull la comparativa de detecció per línia."
)

add_table_caption("Taula 6.6. Detecció anticipada per línia, 6 de Març de 2026.")
add_table(
    ["Línia", "Primer tweet usuari", "Primer avís oficial", "Avantatge", "n pre", "Estat"],
    [
        ["R2N", "07:42", "12:52", "+5 h 10 min",  "4", "Detecció precoç"],
        ["R3",  "11:06", "13:38", "+2 h 32 min",  "1", "Detecció precoç"],
        ["R4",  "06:59", "08:15", "+1 h 16 min",  "1", "Detecció precoç"],
        ["R2",  "13:04", "13:38", "+34 min",       "2", "Detecció precoç"],
        ["R7",  "17:20", "21:54", "+4 h 34 min",   "1", "Detecció precoç"],
        ["R1",  "06:46", "06:51", "+5 min",         "1", "Detecció precoç"],
        ["R8",  "06:53", "21:54", "+15 h 1 min",   "4", "Senyal sense avís immediat"],
        ["R2S", "—",     "06:13", "—",              "0", "Sense detecció prèvia"],
    ]
)

add_para(
    "Per a l'incident de R3 a La Garriga, un tweet publicat a les 11:06 descriu el tall de "
    "via i els retards, amb 2 hores i 32 minuts d'avantatge respecte a la comunicació oficial "
    "de les 13:38. Tot i que la mida mostral és d'un sol tweet pre-oficial, el contingut és "
    "específic i geogràficament precís: menciona La Garriga, Figaró i la línia R3, que "
    "corresponen exactament a la ubicació i l'abast de l'incident. Les línies R2N i R4 "
    "presenten avantatges de 5 hores i 10 minuts i 1 hora i 16 minuts respectivament, tot i "
    "que en aquells casos l'incident de referència no és el robatori de coure. Per a R8, "
    "l'avís oficial molt tardà (21:54) segueix el mateix patró observat al cas del 20 de "
    "Gener: absència de comunicació oficial immediata per a aquella línia durant bona part "
    "de la jornada."
)

# 6.6.5
add_heading("6.6.5 — Síntesi dels quatre casos", level=3)

add_para(
    "La Taula 6.7 recull de forma agregada els principals indicadors de detecció primerenca "
    "per als quatre casos analitzats."
)

add_table_caption("Taula 6.7. Resum dels casos d'estudi.")
add_table(
    ["Data", "Incident principal", "Tweets del dia", "Màx. avantatge", "Línies amb detecció precoç"],
    [
        ["23/10/2025", "Demores i parades",       "506",   "+6 h 51 min (R4)",     "R4, R3"],
        ["20/01/2026", "Temporal + arbre R11",    "807",   "+2 h 21 min (R1/R11)", "R1/R11, R2N, R2S, R2"],
        ["13/02/2026", "Borrasca Nils",           "1.002", "— (event previsible)", "Cap"],
        ["06/03/2026", "Robatori de coure a R3",  "774",   "+5 h 10 min (R2N)",    "R2N, R3, R4, R2, R7, R1"],
    ]
)

add_para(
    "En tres dels quatre casos analitzats, almenys una línia presenta publicacions d'usuaris "
    "classificades com a incident o queixa amb antelació respecte al primer avís oficial. "
    "L'avantatge temporal màxim observat en cadascun d'ells oscil·la entre 33 minuts (R2, "
    "cas de Gener) i 6 hores i 51 minuts (R4, cas d'Octubre). L'única excepció —la Borrasca "
    "Nils del 13 de febrer— no invalida la hipòtesi, sinó que la delimita: el senyal social "
    "té capacitat de detecció anticipada principalment en incidents sobtats i localitzats, "
    "mentre que per als fenòmens meteorològics de gran escala previstos amb antelació, el "
    "canal oficial actua primer [Pender et al., 2014b]. Cal notar que els valors de n "
    "pre-oficial en tots quatre casos es mantenen entre 1 i 6 tweets per línia. Aquesta mida "
    "mostral reduïda és inherent a la tipologia de les incidències de Rodalies —incidents "
    "localitzats que afecten un nombre limitat de viatgers en el moment inicial— i limita "
    "l'aplicació robusta de tests estadístics per línia. El contrast Mann-Whitney U i el delta "
    "de Cliff, inclosos a l'eina de visualització com a validació complementària, s'interpreten "
    "exclusivament en aquells casos en què la mida mostral d'ambdós grups (pre i post oficial) "
    "supera el llindar de cinc observacions."
)

# ── 6.7 ──────────────────────────────────────────────────────────────────────
add_heading("6.7 Resultats de la visualització interactiva", level=2)

add_para(
    "L'eina de visualització desenvolupada (Capítol 5) ha permès integrar les quatre dimensions "
    "del corpus —temporal, geogràfica, per línia i d'incident— en una mateixa interfície "
    "interactiva, facilitant la identificació de patrons que no serien observables des d'una "
    "anàlisi tabular estàtica. El mapa de calor evidencia que les estacions centrals de "
    "Barcelona concentren l'activitat de manera permanent, però les estacions perifèriques "
    "s'activen de forma puntual i precisa en coincidència amb els incidents documentats, la "
    "qual cosa confirma que el corpus conté informació geogràficament rellevant i no tan sols "
    "activitat genèrica sobre el servei."
)
add_para(
    "Els gràfics d'evolució temporal mostren que els dies dels casos d'estudi destaquen "
    "clarament sobre la mediana diària, de manera que l'activitat de Twitter/X actua com a "
    "indicador de la magnitud percebuda de l'incident. Des d'una perspectiva metodològica, "
    "la pàgina de casos d'estudi ha permès verificar de manera sistemàtica i reproductible "
    "la comparativa usuaris/oficial per a cada línia i cada dia. La inclusió del test "
    "estadístic complementari a l'expander de validació (Mann-Whitney U amb correcció "
    "Benjamini-Hochberg i delta de Cliff) ha confirmat, en els pocs casos amb n suficient, "
    "que el classificador manté una confiança estable tant en el període pre-oficial com en "
    "el post-oficial, fet que descarta que la qualitat de la classificació es degradi en les "
    "publicacions inicials —quan la informació és més escassa i menys verificable."
)

# ── 6.8 ──────────────────────────────────────────────────────────────────────
add_heading("6.8 Grau de compliment dels objectius", level=2)

add_para(
    "La Secció 1.3 formula vuit objectius específics per al treball. El primer —la recollida "
    "automatitzada i la preparació d'un corpus de qualitat— s'ha assolit amb un corpus de "
    "66.621 tweets, una pipeline reproducible i una taxa de tweets marcats «en dubte» inferior "
    "a l'1 %. El segon —el disseny d'un sistema de classificació automàtica i l'avaluació del "
    "seu rendiment— s'ha satisfet amb un sistema híbrid que combina regles i model de "
    "llenguatge, documentat a la Secció 4.3, que assoleix una confiança mitjana de 0,8975."
)
add_para(
    "L'extracció d'informació geogràfica i operativa —tercer objectiu— ha produït 12.629 tweets "
    "amb línia assignada, amb una taxa de cobertura del 19,0 % que reflecteix la naturalesa "
    "del corpus: els usuaris rarament especifiquen el codi de línia. La priorització de la "
    "precisió sobre el recall, justificada metodològicament a la Secció 4.2.3, ha permès que "
    "els tweets amb assignació confirmada siguen suficientment fiables per al raonament "
    "comparatiu dels casos d'estudi."
)
add_para(
    "El quart objectiu —l'adaptació de les dades a l'esquema GTFS— s'ha abordat parcialment: "
    "les dades d'estacions s'estructuren internament amb els identificadors de les dades "
    "obertes de l'ATM, però la integració completa amb un feed GTFS operatiu queda com a "
    "línia de treball futur, condicionada a l'establiment d'un marc de col·laboració amb "
    "l'operador."
)

add_note(
    "Revisa si el quart objectiu (GTFS) s'ha implementat més o menys del que s'indica aquí "
    "i ajusta el text en conseqüència."
)

add_para(
    "L'avaluació de la detecció anticipada —cinquè objectiu i nucli empíric del treball— s'ha "
    "completat amb quatre casos d'estudi que demostren avantatges temporals de fins a 6 hores "
    "i 51 minuts en incidents sobtats, i identifiquen com a condició necessària la "
    "imprevisibilitat de l'incident. L'anàlisi de distribució temporal i geogràfica —sisè "
    "objectiu— ha identificat les franges horàries de major risc, la concentració de "
    "l'activitat en el nucli de Barcelona i els pics extraordinaris associats als grans events "
    "d'afectació."
)
add_para(
    "El setè objectiu —el desenvolupament d'una eina de visualització interactiva— s'ha "
    "materialitzat en una aplicació Streamlit de set pàgines que integra les dimensions "
    "espacial, temporal i classificatòria del corpus, tal com es descriu al Capítol 5. "
    "Finalment, la demostració de la viabilitat d'integrar el sistema en plataformes de "
    "monitoratge com Rodalinets o els sistemes SAE dels operadors —vuitè objectiu— s'ha "
    "abordat a nivell de proposta i prova de concepte: els resultats dels casos d'estudi "
    "mostren que el corpus té potencial per complementar els canals de detecció existents, "
    "però la integració operativa en temps real requereix passos addicionals de validació "
    "i acord institucional que s'identifiquen com a treball futur al Capítol 7."
)

# ── Notes de referencies ──────────────────────────────────────────────────────
doc.add_page_break()
add_heading("Notes de referències per al Capítol 6", level=2)
add_note(
    "Afegeix aquestes tres entrades a la Bibliografia general del TFG."
)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
r = p.add_run(
    "[Gu et al., 2016] Gu, Y., Qian, Z. S., & Chen, F. (2016). From Twitter to detector: "
    "Real-time traffic incident detection using social media data. Transportation Research "
    "Part C: Emerging Technologies, 67, 321–342. https://doi.org/10.1016/j.trc.2016.02.011"
)
r.font.name = "Times New Roman"; r.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
r = p.add_run(
    "[Pender et al., 2014a] Pender, B., Currie, G., Delbosc, A., & Shiwakoti, N. (2014). "
    "International study of current and potential social media applications in unplanned "
    "passenger rail disruptions. Transportation Research Record, 2419(1), 118–127. "
    "https://doi.org/10.3141/2419-12"
)
r.font.name = "Times New Roman"; r.font.size = Pt(11)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
r = p.add_run(
    "[Pender et al., 2014b] Pender, B., Currie, G., Delbosc, A., & Shiwakoti, N. (2014). "
    "Social media use during unplanned transit network disruptions: A review of literature. "
    "Transport Reviews, 34(4), 501–521. https://doi.org/10.1080/01441647.2014.915442"
)
r.font.name = "Times New Roman"; r.font.size = Pt(11)

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = r"C:\MARINA\Universitat\TFG - Visualització\Capitol6_Resultats.docx"
doc.save(output_path)
print(f"Document guardat a: {output_path}")
